import docx
import re
import sys
import os

# 检查命令行参数
if len(sys.argv) != 2:
    print('用法: python convert_md_to_docx.py <markdown文件路径>')
    sys.exit(1)

md_file = sys.argv[1]

# 检查文件是否存在
if not os.path.exists(md_file):
    print(f'文件不存在: {md_file}')
    sys.exit(1)

# 生成输出文件名
base_name = os.path.splitext(os.path.basename(md_file))[0]
docx_file = f'{base_name}.docx'

# 读取 Markdown 文件
with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

# 创建 Word 文档
doc = docx.Document()

# 解析 Markdown 内容
lines = content.split('\n')

current_style = None
current_heading_level = 0

for line in lines:
    line = line.strip()
    
    # 处理标题
    if line.startswith('#'):
        heading_level = line.count('#')
        title = line.strip('#').strip()
        
        if heading_level == 1:
            doc.add_heading(title, level=0)
        elif heading_level == 2:
            doc.add_heading(title, level=1)
        elif heading_level == 3:
            doc.add_heading(title, level=2)
        
        current_heading_level = heading_level
        current_style = 'heading'
    
    # 处理目录项
    elif line.startswith('- ['):
        # 提取链接文本
        match = re.match(r'- \[(.*?)\]\(.*?\)', line)
        if match:
            text = match.group(1)
            doc.add_paragraph(text, style='List Bullet')
        current_style = 'list'
    
    # 处理列表项
    elif line.startswith('- A.') or line.startswith('- B.') or line.startswith('- C.') or line.startswith('- D.'):
        doc.add_paragraph(line, style='List Bullet')
        current_style = 'list'
    
    # 处理正确答案
    elif line.startswith('**正确答案：'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        current_style = 'answer'
    
    # 处理解释
    elif line.startswith('**解释**：'):
        p = doc.add_paragraph()
        # 提取粗体部分和普通文本
        if '**解释**：' in line:
            p.add_run('**解释**：').bold = True
            p.add_run(line.replace('**解释**：', ''))
        current_style = 'explanation'
    
    # 处理分隔线
    elif line == '---':
        doc.add_paragraph('')
        current_style = 'separator'
    
    # 处理空行
    elif line == '':
        if current_style not in ['separator', 'heading']:
            doc.add_paragraph('')
    
    # 处理普通文本
    else:
        doc.add_paragraph(line, style='Normal')
        current_style = 'normal'

# 保存 Word 文档
doc.save(docx_file)
print(f'转换完成！Word 文档已保存为 {docx_file}')
